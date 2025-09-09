import os, sys
import pandas as pd
from docx import Document as DocxDocument
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.schema import Document

def load_documents(logger=None, FILE_DIR=None):
    """
    Carica i documenti da un file (PDF o TXT), li splitta in chunk
    e aggiunge i metadata (source = nome file).
    """
    filename = os.path.basename(FILE_DIR)

    if logger: 
        logger.info("--------")
        logger.info("Caricamento e suddivisione documenti...")

    if not os.path.exists(FILE_DIR):
        if logger: logger.error(f"❌ File non trovato: {FILE_DIR}")
        sys.exit(1)

    ext = os.path.splitext(FILE_DIR)[1].lower()
    splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)

    try:
        if ext == ".pdf":
            if logger: logger.info("📄 Caricamento documento PDF...")
            loader = PyPDFLoader(FILE_DIR)
            pages = loader.load()
            docs = splitter.split_documents(pages)

        elif ext == ".txt":
            if logger: logger.info("📄 Caricamento documento TXT...")
            with open(FILE_DIR, "r", encoding="utf-8") as f:
                text = f.read()
            docs = splitter.split_documents([Document(page_content=text, metadata={"source": FILE_DIR})])

        elif ext == ".docx":
            if logger: logger.info("📄 Caricamento documento DOCX...")
            doc = DocxDocument(FILE_DIR)
            full_text = "\n".join([p.text for p in doc.paragraphs])
            docs = splitter.split_documents([Document(page_content=full_text, metadata={"source": FILE_DIR})])

        elif ext == ".csv":
            if logger: logger.info("📄 Caricamento documento CSV...")
            df = pd.read_csv(FILE_DIR)
            text = df.to_string()
            docs = splitter.split_documents([Document(page_content=text, metadata={"source": FILE_DIR})])

        else:
            raise ValueError(f"❌ Formato non supportato: {ext}")
         

        #Aggiungo il campo "source" ai metadata di ciascun chunk
        for d in docs:
            d.metadata["source"] = filename


        logger.info(f"📑 Estratti {len(docs)} chunk da {filename} (source impostato)")
        return docs


    except Exception as e:
        if logger: logger.error(f"❌ Errore durante il caricamento di {FILE_DIR}: {str(e)}")
        raise

